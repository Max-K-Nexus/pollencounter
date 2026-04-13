#!/usr/bin/env python3
"""
Script per automatizzare la conta pollinica.
Versione con foglio riepilogo settimanale.

Ogni codice inserito:
  - viene registrato nel foglio 'dati_grezzi' (log con timestamp)
  - incrementa di +1 la cella corrispondente nel foglio 'riepilogo_settimana'
  - la tabella concentrazione si aggiorna automaticamente (formule)

Il template Polline_Template_Settimanale.xlsx NON viene modificato.
"""

import json
import re
import signal
import sys
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path

try:
    import winsound
except ImportError:
    winsound = None

try:
    import docx as _docx_module
except ImportError:
    _docx_module = None

try:
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    print("ERRORE: openpyxl non installato. Installa con:")
    print("  pip3 install openpyxl")
    sys.exit(1)

# Supporto PyInstaller
if getattr(sys, "frozen", False):
    BUNDLE_DIR = Path(sys._MEIPASS)
    OUTPUT_DIR = Path(sys.executable).parent
else:
    BUNDLE_DIR = Path(__file__).parent
    OUTPUT_DIR = Path(__file__).parent

TEMPLATE_FILE = BUNDLE_DIR / "Polline_Template_Settimanale.xlsx"

AUTOSAVE_INTERVAL = 5  # salva ogni N inserimenti

# ── Stili openpyxl riutilizzati in piu' funzioni ──
THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)
FILL_GIALLO = PatternFill("solid", fgColor="FFE699")
FILL_VERDE = PatternFill("solid", fgColor="92D050")
FILL_VERDE_CHIARO = PatternFill("solid", fgColor="C5E0B4")
FILL_BLU = PatternFill("solid", fgColor="4472C4")
FONT_BOLD = Font(bold=True)
FONT_BOLD_BIG = Font(bold=True, size=14)
FONT_BIANCO_BOLD = Font(color="FFFFFF", bold=True)
ALIGN_CENTER = Alignment(horizontal="center")
ALIGN_CENTER_WRAP = Alignment(horizontal="center", wrap_text=True)


def _beep():
    """Emette un beep sonoro (winsound su Windows, bell ASCII come fallback)."""
    if winsound is not None:
        winsound.Beep(800, 150)
    else:
        print("\a", end="", flush=True)


GIORNI_NOMI = {
    1: "lunedi", 2: "martedi", 3: "mercoledi",
    4: "giovedi", 5: "venerdi", 6: "sabato", 7: "domenica",
}


CODICI_SPECIE = {
    "01": "ACERACEAE", "02": "ALTRI POLLINI", "03": "BETULACEAE",
    "04": "Alnus", "05": "Betula", "06": "CANNABACEAE",
    "07": "CHENO-AMAR", "08": "COMPOSITAE", "09": "Altre compositae",
    "10": "Ambrosia", "11": "Artemisia", "12": "CORYLACEAE (somma c+o)",
    "13": "Carpinus/Ostrya", "14": "Carpinus", "15": "Ostrya carpinifolia",
    "16": "Corylus avellana", "17": "CUP-TAXACEAE", "18": "ERICACEAE",
    "19": "EUPHORBIACEAE", "20": "FAGACEAE", "21": "Castanea sativa",
    "22": "Fagus sylvatica", "23": "Quercus", "24": "GRAMINEAE",
    "25": "HIPPOCASTANACEAE", "26": "JUGLANDACEAE", "27": "LAURACEAE",
    "28": "MIMOSACEAE", "29": "MORACEAE", "30": "MYRTACEAE",
    "31": "OLEACEAE", "32": "Altre oleaceae", "33": "Fraxinus",
    "34": "Ligustrum", "35": "Olea", "36": "PINACEAE",
    "37": "PLANTAGINACEAE", "38": "PLATANACEAE",
    "39": "POLLINI NON IDENTIFICATI", "40": "POLYGONACEAE",
    "41": "SALICACEAE", "42": "Populus", "43": "Salix",
    "44": "TILIACEAE", "45": "ULMACEAE", "46": "UMBELLIFERAE",
    "47": "URTICACEAE",
    "48": "Alternaria", "49": "Botrytis", "50": "Cladosporium",
    "51": "Curvularia", "52": "Epicoccum", "53": "Helminthosporium",
    "54": "Pithomyces", "55": "Pleospora", "56": "Polythrincium",
    "57": "Stemphylium", "58": "Tetraploa", "59": "Torula",
}

# Mapping codice specie -> nome famiglia nel file soglie
SOGLIE_MAPPING = {
    "01": "Aceracee",
    "03": "Betulaceae",
    "04": "Betulaceae",
    "05": "Betulaceae",
    "07": "Cheno-Amarantaceae",
    "08": "Composite",
    "09": "Composite",
    "10": "Composite",
    "11": "Composite",
    "12": "Corilacee",
    "13": "Corilacee",
    "14": "Corilacee",
    "15": "Corilacee",
    "16": "Corilacee",
    "17": "Cupressaceae + Taxaceae",
    "20": "Fagaceae",
    "21": "Fagaceae",
    "22": "Fagaceae",
    "23": "Fagaceae",
    "24": "Graminaceae",
    "31": "Oleaceae",
    "32": "Oleaceae",
    "33": "Oleaceae",
    "34": "Oleaceae",
    "35": "Oleaceae",
    "36": "Pinaceae",
    "37": "Plantaginaceae",
    "38": "Platanaceae",
    "41": "Salicaceae",
    "42": "Salicaceae",
    "43": "Salicaceae",
    "45": "Ulmaceae",
    "47": "Urticaceae",
    "48": "Alternaria",
    "50": "Cladosporium",
}

# ── Costanti riepilogo annuale ──
POLLINI_CODICI = [f"{i:02d}" for i in range(1, 48)]    # 47 pollini
SPORE_CODICI = [f"{i:02d}" for i in range(48, 60)]      # 12 spore

ANNUALE_VERDE_CODICI = {
    "05", "12", "13", "14", "15", "17", "18", "19",
    "22", "23", "24", "33", "35", "36", "47",
}
ANNUALE_VERDE_CHIARO_CODICI = {"48"}
ANNUALE_BOLD_CODICI = ANNUALE_VERDE_CODICI | ANNUALE_VERDE_CHIARO_CODICI

# Layout colonne riepilogo annuale (conta grezza)
_ANN_POLL_START = 2                                           # col B
_ANN_SEP1 = _ANN_POLL_START + len(POLLINI_CODICI)            # 49
_ANN_SPORE_START = _ANN_SEP1 + 1                             # 50
_ANN_SPORE_END = _ANN_SPORE_START + len(SPORE_CODICI) - 1   # 61
# Gap
_ANN_GAP = _ANN_SPORE_END + 1                                # 62
# Layout colonne riepilogo annuale (concentrazioni)
_ANN_CONC_DATA = _ANN_GAP + 1                                # 63
_ANN_CONC_POLL_START = _ANN_CONC_DATA + 1                    # 64
_ANN_CONC_SEP = _ANN_CONC_POLL_START + len(POLLINI_CODICI)   # 111
_ANN_CONC_SPORE_START = _ANN_CONC_SEP + 1                    # 112
_ANN_CONC_SPORE_END = _ANN_CONC_SPORE_START + len(SPORE_CODICI) - 1  # 123

# ── Costanti bollettino Word ──
_MESI_ITA = ["Gennaio", "Febbraio", "Marzo", "Aprile", "Maggio", "Giugno",
             "Luglio", "Agosto", "Settembre", "Ottobre", "Novembre", "Dicembre"]
_GIORNI_ITA_LONG = ["Luned\xec", "Marted\xec", "Mercoled\xec", "Gioved\xec", "Venerd\xec", "Sabato", "Domenica"]
_GIORNI_ENG_LONG = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

# (nome_ita, nome_eng, riep_rows, soglie_famiglia, fallback_assente, fallback_bassa, fallback_media)
# riep_rows: righe nel foglio riepilogo_settimana da sommare (vedi codice_to_row)
# soglie_famiglia: nome famiglia nel file soglie (sovrascrive i valori fallback se presente)
BOLL_WORD_RIGHE = [
    ("ACERACEAE",           "ACERACEAE",              [6],                  "Aceracee",                0.9,   19.9,  39.9),
    ("BETULACEAE",          "BETULACEAE",              [8, 9, 10],           "Betulaceae",              0.5,   15.9,  49.9),
    ("Alnus",               "Alnus",                  [9],                  "Betulaceae",              0.5,   15.9,  49.9),
    ("Betula",              "Betula",                 [10],                  "Betulaceae",              0.5,   15.9,  49.9),
    ("CHENO-AMAR",          "CHENO-AMAR",             [12],                  "Cheno-Amarantaceae",      0.0,    4.9,  24.9),
    ("COMPOSITAE",          "COMPOSITAE",             [13, 14, 15, 16],      "Composite",               0.0,    4.9,  24.9),
    ("Altre compositae",    "Other composites",       [14],                  "Composite",               0.0,    4.9,  24.9),
    ("Ambrosia",            "Ragweed",                [15],                  "Composite",               0.0,    4.9,  24.9),
    ("Artemisia",           "Mugwort",                [16],                  "Composite",               0.0,    4.9,  24.9),
    ("CORYLACEAE",          "CORYLACEAE",             [17, 18, 19, 20, 21],  "Corilacee",               0.5,   15.9,  49.9),
    ("Carpinus/Ostrya",     "Hornbeam/Hop hornbeam",  [18],                  "Corilacee",               0.5,   15.9,  49.9),
    ("Carpinus",            "Hornbeam",               [19],                  "Corilacee",               0.5,   15.9,  49.9),
    ("Ostrya carpinifolia", "Hop hornbeam",           [20],                  "Corilacee",               0.5,   15.9,  49.9),
    ("Corylus avellana",    "Hazel",                  [21],                  "Corilacee",               0.5,   15.9,  49.9),
    ("CUP-TAXACEAE",        "Cup-Taxaceae",           [22],                  "Cupressaceae + Taxaceae", 3.9,   29.9,  89.9),
    ("FAGACEAE",            "FAGACEAE",               [25, 26, 27, 28],      "Fagaceae",                0.9,   19.9,  39.9),
    ("Castanea sativa",     "Sweet chestnut",         [26],                  "Fagaceae",                0.9,   19.9,  39.9),
    ("Fagus sylvatica",     "Beech",                  [27],                  "Fagaceae",                0.9,   19.9,  39.9),
    ("Quercus",             "Oak",                    [28],                  "Fagaceae",                0.9,   19.9,  39.9),
    ("GRAMINEAE",           "Grasses",                [29],                  "Graminaceae",             0.5,    9.9,  29.9),
    ("OLEACEAE",            "OLEACEAE",               [36, 37, 38, 39, 40],  "Oleaceae",                0.5,    4.9,  24.9),
    ("Altre oleaceae",      "Other olive family",     [37],                  "Oleaceae",                0.5,    4.9,  24.9),
    ("Fraxinus",            "Ash",                    [38],                  "Oleaceae",                0.5,    4.9,  24.9),
    ("Ligustrum",           "Privet",                 [39],                  "Oleaceae",                0.5,    4.9,  24.9),
    ("Olea",                "Olive",                  [40],                  "Oleaceae",                0.5,    4.9,  24.9),
    ("PINACEAE",            "PINACEAE",               [41],                  "Pinaceae",                0.9,   14.9,  49.9),
    ("PLANTAGINACEAE",      "PLANTAGINACEAE",         [42],                  "Plantaginaceae",          0.0,    0.4,   1.9),
    ("PLATANACEAE",         "Plane tree",             [43],                  "Platanaceae",             0.9,   19.9,  39.9),
    ("SALICACEAE",          "SALICACEAE",             [46, 47, 48],          "Salicaceae",              0.9,   19.9,  39.9),
    ("Populus",             "Poplar",                 [47],                  "Salicaceae",              0.9,   19.9,  39.9),
    ("Salix",               "Willow",                 [48],                  "Salicaceae",              0.9,   19.9,  39.9),
    ("ULMACEAE",            "Elm family",             [50],                  "Ulmaceae",                0.9,   19.9,  39.9),
    ("URTICACEAE",          "Nettle",                 [52],                  "Urticaceae",              1.9,   19.9,  69.9),
    ("Alternaria",          "Alternaria",             [58],                  "Alternaria",              1.9,   19.0, 100.0),
    ("Cladosporium",        "Cladosporium",           [60],                  "Cladosporium",          100.0,  499.0,1000.0),
]

# Layout foglio Calendario (specie in righe, date in colonne)
_CAL_HEADER_ROW = 3       # riga date intestazione
_CAL_DOY_ROW = 4          # riga giorno-dell'anno
_CAL_DATA_START_ROW = 5   # prima riga dati specie
_CAL_COL_SPECIE = 1
_CAL_COL_DATE_START = 2
_CAL_SEP_ROW = _CAL_DATA_START_ROW + len(POLLINI_CODICI)  # 52

SCRIPT_DIR = Path(__file__).parent if not getattr(sys, "frozen", False) else Path(sys._MEIPASS)

# Cartella dove risiede il file di configurazione (accanto allo script o all'exe)
_CONFIG_DIR = Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).parent
CONFIG_FILE = _CONFIG_DIR / "pollencounter.cfg"

# True quando lo script e' avviato come subprocess dalla GUI (--gui in argv)
_GUI_MODE = "--gui" in sys.argv



# ============================================================
# Helper di basso livello
# ============================================================
def codice_to_row(codice):
    """Codice specie -> riga nel foglio riepilogo_settimana."""
    n = int(codice)
    if 1 <= n <= 47:          # Pollini
        return n + 5
    if 48 <= n <= 59:         # Spore
        return n + 10
    return None


def giorno_to_col(giorno_num):
    """Giorno (1=lun, 7=dom) -> colonna dati grezzi (G=7 ... M=13)."""
    return giorno_num + 6


def normalizza_codice(codice):
    """Normalizza un codice a 2 cifre (es. '5' -> '05')."""
    if codice.isdigit() and len(codice) == 1:
        return codice.zfill(2)
    return codice


def leggi_valore(ws, row, col):
    """Legge un valore intero da una cella. Ritorna 0 se vuota o non numerica."""
    val = ws.cell(row=row, column=col).value
    if isinstance(val, (int, float)):
        return int(val)
    return 0


def leggi_fattore(ws_riepilogo):
    """Legge il fattore di conversione dalla cella Q3 del foglio riepilogo.
    Ritorna 0.4 come default se la cella e' vuota o non valida."""
    val = ws_riepilogo["Q3"].value
    if isinstance(val, (int, float)) and val > 0:
        return float(val)
    return 0.4


def scrivi_log(ws_log, log_row, data_str, codice, specie, nota=None):
    """Scrive una riga nel foglio dati_grezzi con timestamp automatico."""
    ws_log.cell(row=log_row, column=1, value=data_str)
    ws_log.cell(row=log_row, column=2, value=codice)
    ws_log.cell(row=log_row, column=3, value=specie)
    ws_log.cell(row=log_row, column=4, value=datetime.now().strftime("%H:%M:%S"))
    if nota:
        ws_log.cell(row=log_row, column=5, value=nota)


def cancella_riga_log(ws_log, log_row):
    """Svuota una riga del log (colonne 1-5)."""
    for c in range(1, 6):
        ws_log.cell(row=log_row, column=c, value=None)


# ============================================================
# Funzioni di utilita'
# ============================================================
def display_menu():
    print("\n" + "=" * 60)
    print("CONTA POLLINICA - SISTEMA AUTOMATIZZATO")
    print("=" * 60)
    print("\nCodici disponibili:\n")
    keys = list(CODICI_SPECIE.keys())
    for i in range(0, len(keys), 2):
        code1, specie1 = keys[i], CODICI_SPECIE[keys[i]]
        line = f"  {code1}: {specie1:<30}"
        if i + 1 < len(keys):
            code2, specie2 = keys[i + 1], CODICI_SPECIE[keys[i + 1]]
            line += f"  {code2}: {specie2}"
        print(line)
    print("\n" + "-" * 60)
    print("Comandi:")
    print("  01-59   Inserisce la specie corrispondente")
    print("  NNxQ    Inserisce Q occorrenze (es. 48x4 = 4 Alternaria)")
    print("  .       Ripete l'ultimo codice inserito")
    if not _GUI_MODE:
        print("  r       Riepilogo giornata corrente")
        print("  w       Riepilogo settimanale (tutti i giorni)")
    print("  l       Ultimi inserimenti (storico)")
    print("  c       Correggi un giorno precedente")
    print("  n       Aggiungi una nota per la giornata")
    print("  u       Annulla ultimo inserimento")
    print("  b       Attiva/disattiva beep sonoro")
    print("  h       Mostra questo menu")
    print("  s       Salva il file (senza uscire)")
    print("  d       Chiudi giornata (puoi continuare con un altro giorno)")
    print("  q       Salva il file e esci")
    print("-" * 60 + "\n")


def parse_data_flessibile(testo):
    """Cerca di estrarre una data da testo libero.

    Formati riconosciuti (giorno e mese anche a 1 cifra):
      9-2-2026    9/2/2026    09-02-2026    09/02/2026
      9-2-26      9/2/26      09-02-26      09/02/26
      9 febbraio 2026   9 feb 2026   (mesi italiani)
    Ritorna un datetime oppure None.
    """
    testo = testo.strip()

    # 1) Formato numerico: G-M-A oppure G/M/A (separatore - o /)
    match = re.search(r"(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})", testo)
    if match:
        g, m, a = match.group(1), match.group(2), match.group(3)
        if len(a) == 2:
            a = "20" + a
        try:
            return datetime(int(a), int(m), int(g))
        except ValueError:
            pass

    # 2) Formato testuale: "9 febbraio 2026" o "9 feb 2026"
    mesi_txt = {
        "gen": 1, "gennaio": 1, "feb": 2, "febbraio": 2,
        "mar": 3, "marzo": 3, "apr": 4, "aprile": 4,
        "mag": 5, "maggio": 5, "giu": 6, "giugno": 6,
        "lug": 7, "luglio": 7, "ago": 8, "agosto": 8,
        "set": 9, "settembre": 9, "ott": 10, "ottobre": 10,
        "nov": 11, "novembre": 11, "dic": 12, "dicembre": 12,
    }
    match = re.search(r"(\d{1,2})\s+([a-zA-Z]+)\s+(\d{2,4})", testo)
    if match:
        g = int(match.group(1))
        m_txt = match.group(2).lower()
        a = match.group(3)
        if len(a) == 2:
            a = "20" + a
        m = mesi_txt.get(m_txt)
        if m:
            try:
                return datetime(int(a), m, g)
            except ValueError:
                pass

    return None


def leggi_settimana_da_file(ws):
    """Legge la data del lunedi' dalla cella J3 del foglio riepilogo.
    Ritorna un datetime oppure None se la cella e' vuota o non leggibile."""
    val = ws["J3"].value
    if not val:
        return None
    if isinstance(val, datetime):
        return val
    if isinstance(val, str):
        return parse_data_flessibile(val)
    return None


def chiedi_settimana(settimana_file=None):
    """Chiede l'intervallo settimanale. Ritorna la data del lunedi' (datetime).

    Se settimana_file e' fornita (datetime), propone di mantenerla.
    """
    oggi = datetime.now()
    lun_corrente = oggi - timedelta(days=oggi.weekday())
    dom_corrente = lun_corrente + timedelta(days=6)

    # Se il file contiene gia' una settimana, proponi di mantenerla
    if settimana_file:
        lun_file = settimana_file - timedelta(days=settimana_file.weekday())
        dom_file = lun_file + timedelta(days=6)
        print(f"\nIl file si riferisce alla settimana:")
        print(f"  dal {lun_file.strftime('%d-%m-%Y')} al {dom_file.strftime('%d-%m-%Y')}")
        risp = input("  Mantenere questa settimana? (s/n): ").strip().lower()
        if risp != "n":
            print(f"  -> Settimana: {lun_file.strftime('%d-%m-%Y')} (lun) "
                  f"- {dom_file.strftime('%d-%m-%Y')} (dom)")
            return lun_file

    default_str = (f"dal {lun_corrente.strftime('%d-%m-%Y')} "
                   f"al {dom_corrente.strftime('%d-%m-%Y')}")

    print(f"\nChe settimana stiamo analizzando?")
    print(f"  Invio = settimana corrente ({default_str})")
    print(f"  Accetta: 9-2-2026  9/2/2026  09-02-26  9 feb 2026  ecc.")
    while True:
        inp = input("  Settimana: ").strip()
        if not inp:
            lunedi = lun_corrente
        else:
            dt = parse_data_flessibile(inp)
            if not dt:
                print("  Data non riconosciuta. Prova con: 9-2-2026 o 9/2/2026 o 9 feb 2026")
                continue
            lunedi = dt - timedelta(days=dt.weekday())
        domenica = lunedi + timedelta(days=6)
        print(f"  -> Settimana: {lunedi.strftime('%d-%m-%Y')} (lun) "
              f"- {domenica.strftime('%d-%m-%Y')} (dom)")
        return lunedi


def chiedi_giorno(lunedi):
    """Chiede il giorno della settimana. Ritorna (giorno_num, data_str)."""
    print("\nSeleziona il giorno di lavoro:")
    for num, nome in GIORNI_NOMI.items():
        data_giorno = lunedi + timedelta(days=num - 1)
        print(f"  {num}) {nome.upper():<12} {data_giorno.strftime('%d-%m-%Y')}")
    while True:
        scelta = input("Scegli (1-7): ").strip()
        if scelta in [str(x) for x in range(1, 8)]:
            giorno_num = int(scelta)
            data_giorno = lunedi + timedelta(days=giorno_num - 1)
            return giorno_num, data_giorno.strftime("%d-%m-%Y")
        print("Scelta non valida, riprova.")


def compila_intestazione(ws, lunedi):
    """Compila i metadati della settimana nel foglio riepilogo (riga 3)."""
    domenica = lunedi + timedelta(days=6)
    mese_nome = _MESI_ITA[lunedi.month - 1]
    anno = lunedi.year
    fmt = "%d-%m-%Y"

    # Lato sinistro (grezzi)
    ws["H3"] = mese_nome
    ws["J3"] = lunedi          # data Excel: le formule del bollettino la referenziano
    ws["J3"].number_format = "DD-MM-YYYY"
    ws["K3"] = domenica
    ws["K3"].number_format = "DD-MM-YYYY"
    ws["M3"] = anno

    # Lato destro (concentrazione) — stessi valori
    ws["T3"] = mese_nome
    ws["V3"] = lunedi.strftime(fmt)
    ws["W3"] = domenica.strftime(fmt)
    ws["Y3"] = anno


def find_next_log_row(ws):
    for row in range(2, ws.max_row + 2):
        if ws.cell(row=row, column=1).value is None:
            return row
    return ws.max_row + 1


def chiedi_percorso_salvataggio(nome_default):
    """Chiede il percorso di salvataggio del file.

    GUI: il marker apre asksaveasfilename; la risposta e' il percorso completo.
    CLI: l'utente digita il nome (o invio per il default) nella cartella corrente.
    Ritorna un Path, oppure None se l'utente ha annullato dalla dialog GUI.
    """
    print(f"\n  File: {nome_default}")
    print(f"__GUI_ASKSAVEFILE__|{OUTPUT_DIR}|{nome_default}", flush=True)
    risposta = input("  Percorso (invio = nome predefinito nella cartella corrente): ").strip()
    if not risposta:
        if _GUI_MODE:
            return None          # GUI: dialog annullata
        return OUTPUT_DIR / nome_default   # CLI: usa il default
    p = Path(risposta)
    if p.suffix.lower() != ".xlsx":
        p = p.with_suffix(".xlsx")
    if not p.is_absolute():
        p = OUTPUT_DIR / p
    return p


def menu_uscita_salvataggio(wb, prima_data, nome_ripreso, file_ripreso, stato=None):
    """Chiede se salvare e dove. Ritorna (True, cartella) o (False, None)."""
    # Nome predefinito: file ripreso se disponibile, altrimenti Conta_Pollinica_<data>
    data_pulita = re.sub(r"^(Conta_Pollinica_|~autosave_|incompleto_)+", "", prima_data)
    nome_default = nome_ripreso or f"Conta_Pollinica_{data_pulita or prima_data}.xlsx"
    if not nome_default.endswith(".xlsx"):
        nome_default += ".xlsx"

    # Se gia' salvato mid-session, proponi salvataggio rapido sullo stesso percorso
    percorso_precedente = stato.get("percorso_salvato") if stato else None

    print("\n" + "=" * 60)
    while True:
        risp = input("\nSalvare e uscire? (s/n): ").strip().lower()
        if risp == "s":
            percorso = None
            # Salvataggio rapido: file gia' salvato mid-session o file ripreso
            if percorso_precedente:
                risp_rapido = input(f"  Salvare su '{percorso_precedente.name}'? (s = si, n = altro percorso): ").strip().lower()
                if risp_rapido == "s":
                    percorso = percorso_precedente
            elif file_ripreso and Path(file_ripreso).exists():
                risp_rapido = input(f"  Salvare su '{Path(file_ripreso).name}'? (s = si, n = altro percorso): ").strip().lower()
                if risp_rapido == "s":
                    percorso = Path(file_ripreso)

            if percorso is None:
                percorso = chiedi_percorso_salvataggio(nome_default)
                if percorso is None:
                    continue   # GUI: dialog annullata, riproponi
                if percorso.exists() and not _GUI_MODE:
                    risp2 = input(f"  '{percorso.name}' esiste gia'. Sovrascrivere? (s/n): ").strip().lower()
                    if risp2 != "s":
                        continue

            _attendi_autosave()
            wb.save(percorso)
            print(f"\n  [OK] File salvato: {percorso}")
            return True, percorso.parent
        elif risp == "n":
            risp2 = input("Uscire senza salvare? I dati non salvati andranno persi. (s/n): ").strip().lower()
            if risp2 == "s":
                print("\n  [OK] Uscita senza salvataggio.")
                return False, None


def cerca_file_esistenti(cartella=None):
    """Cerca file .xlsx in cartella (default OUTPUT_DIR) escludendo il template.
    Ritorna lista di Path ordinata per data modifica (piu' recente prima).
    """
    if cartella is None:
        cartella = OUTPUT_DIR
    template_name = TEMPLATE_FILE.name
    files = [f for f in cartella.glob("*.xlsx")
             if f.name != template_name and f.is_file()]
    files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    return files


def _conta_righe_log(path):
    """Ritorna una stringa '(N righe log)' per il file .xlsx indicato."""
    try:
        wb_tmp = openpyxl.load_workbook(path, read_only=True)
        ws_tmp = wb_tmp["dati_grezzi"]
        n = sum(1 for row in ws_tmp.iter_rows(min_row=2, max_col=1)
                if row[0].value is not None)
        wb_tmp.close()
        return f"({n} righe log)"
    except Exception:
        return ""



def chiedi_ripresa_o_nuovo(file_esistenti):
    """Mostra menu di ripresa. Ritorna Path scelto oppure None (= nuovo)."""
    print("\nFile esistenti trovati:")
    for i, f in enumerate(file_esistenti, 1):
        print(f"  {i}) {f.name}  {_conta_righe_log(f)}")
    print(f"  n) Nuovo file (dal template)")
    print(f"  i) Importa da altra cartella")

    while True:
        scelta = input("Scelta: ").strip().lower()
        if scelta == "n":
            return None
        if scelta == "i":
            print(f"__GUI_ASKOPENFILE__|{OUTPUT_DIR}", flush=True)
            raw = input("  Percorso file (invio per annullare): ").strip()
            if not raw:
                continue
            p = Path(raw)
            if p.is_file() and p.suffix.lower() == ".xlsx":
                return p
            elif p.is_dir():
                files_import = cerca_file_esistenti(p)
                if not files_import:
                    print("  Nessun file .xlsx trovato in quella cartella.")
                    continue
                print(f"\n  File trovati in {p}:")
                for j, f in enumerate(files_import, 1):
                    print(f"    {j}) {f.name}  {_conta_righe_log(f)}")
                while True:
                    sc = input("  Scelta (numero, invio per annullare): ").strip()
                    if not sc:
                        break
                    if sc.isdigit():
                        idx2 = int(sc) - 1
                        if 0 <= idx2 < len(files_import):
                            return files_import[idx2]
                    print("  Scelta non valida.")
            else:
                print("  Percorso non valido.")
            continue
        if scelta.isdigit():
            idx = int(scelta) - 1
            if 0 <= idx < len(file_esistenti):
                return file_esistenti[idx]
        print("Scelta non valida, riprova.")


# ============================================================
# Bollettino pollinico
# ============================================================
def _parse_soglia_max(testo):
    """Estrae il valore massimo da un range come '0 - 0,5', '< 1', '> 50'.
    Ritorna un float."""
    if not testo:
        return 0.0
    testo = str(testo).strip()
    # Sostituisci virgole decimali
    testo = testo.replace(",", ".")
    # Formato "> N" — non ha max, ritorna infinito
    if testo.startswith(">"):
        return float("inf")
    # Formato "< N" — N è escluso, usiamo N-0.1 come soglia superiore dell'intervallo
    m = re.match(r"<\s*([\d.]+)", testo)
    if m:
        return float(m.group(1)) - 0.1
    # Formato "N - M" — prendi M
    m = re.search(r"([\d.]+)\s*[-–]\s*([\d.]+)", testo)
    if m:
        return float(m.group(2))
    # Numero singolo
    m = re.match(r"([\d.]+)", testo)
    if m:
        return float(m.group(1))
    return 0.0


def _parse_soglie_da_foglio(ws):
    """Parsing comune: legge soglie da un foglio Excel (formato concentrazioni).
    Ritorna dict {nome_famiglia: (max_assente, max_bassa, max_media)}."""
    soglie = {}
    for row in ws.iter_rows(min_row=3, max_col=5):
        nome = row[0].value
        if not nome or not isinstance(nome, str):
            continue
        nome = nome.strip()
        # Salta righe titolo/sezione senza valori numerici (es. "Spore Fungine")
        if not row[1].value or not any(c.isdigit() for c in str(row[1].value)):
            continue
        max_assente = _parse_soglia_max(row[1].value)
        max_bassa = _parse_soglia_max(row[2].value)
        max_media = _parse_soglia_max(row[3].value)
        soglie[nome] = (max_assente, max_bassa, max_media)
    return soglie


def carica_soglie(wb=None):
    """Carica le soglie di concentrazione pollinica.
    Se wb e' fornito e contiene un foglio 'soglie', legge da li'.
    Altrimenti fallback al file esterno concentrazioni_polliniche.xlsx.
    Ritorna dict {nome_famiglia: (max_assente, max_bassa, max_media)}."""

    # Tentativo 1: foglio "soglie" nel workbook corrente
    if wb is not None and "soglie" in wb.sheetnames:
        return _parse_soglie_da_foglio(wb["soglie"])

    # Tentativo 2: file esterno (fallback)
    nome_file = "concentrazioni_polliniche.xlsx"
    candidati = [
        OUTPUT_DIR / nome_file,
        SCRIPT_DIR / nome_file,
        SCRIPT_DIR.parent / nome_file,
        BUNDLE_DIR / nome_file,
    ]
    filepath = None
    for c in candidati:
        if c.exists():
            filepath = c
            break
    if filepath is None:
        print(f"  ERRORE: file soglie '{nome_file}' non trovato.")
        print(f"  Cercato in: {', '.join(str(c) for c in candidati)}")
        return None

    wb_soglie = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    soglie = _parse_soglie_da_foglio(wb_soglie.active)
    wb_soglie.close()
    return soglie


# ============================================================
# Riepilogo annuale
# ============================================================
def raccogli_dati_giornalieri(ws_riepilogo):
    """Raccoglie dati dal foglio riepilogo settimanale.
    Ritorna {giorno_num: {codice: valore}} solo per giorni con dati."""
    risultato = {}
    tutti_codici = POLLINI_CODICI + SPORE_CODICI
    for giorno_num in range(1, 8):
        col = giorno_to_col(giorno_num)
        dati = {}
        for codice in tutti_codici:
            row = codice_to_row(codice)
            if row is None:
                continue
            val = leggi_valore(ws_riepilogo, row, col)
            if val > 0:
                dati[codice] = val
        if dati:
            risultato[giorno_num] = dati
    return risultato


def _ann_col_grezzo(codice):
    """Colonna conta grezza nel riepilogo annuale per un codice."""
    n = int(codice)
    if 1 <= n <= 47:
        return _ANN_POLL_START + (n - 1)
    if 48 <= n <= 59:
        return _ANN_SPORE_START + (n - 48)
    return None


def _ann_col_conc(codice):
    """Colonna concentrazione nel riepilogo annuale per un codice."""
    n = int(codice)
    if 1 <= n <= 47:
        return _ANN_CONC_POLL_START + (n - 1)
    if 48 <= n <= 59:
        return _ANN_CONC_SPORE_START + (n - 48)
    return None


def crea_intestazione_annuale(ws, anno):
    """Crea le righe 1-3 del riepilogo annuale (titolo, sezioni, intestazioni)."""
    # Riga 1: titolo
    cell = ws.cell(row=1, column=1, value=f"RIEPILOGO ANNUALE {anno}")
    cell.font = FONT_BOLD_BIG
    cell.alignment = ALIGN_CENTER
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=_ANN_SPORE_END)

    # Riga 2: sezioni
    cell = ws.cell(row=2, column=1, value="CONTA GREZZA")
    cell.font = FONT_BOLD
    cell.alignment = ALIGN_CENTER
    cell.fill = FILL_GIALLO
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=_ANN_SPORE_END)

    cell = ws.cell(row=2, column=_ANN_CONC_DATA, value="CONCENTRAZIONI (p/m3)")
    cell.font = FONT_BOLD
    cell.alignment = ALIGN_CENTER
    cell.fill = FILL_GIALLO
    ws.merge_cells(start_row=2, start_column=_ANN_CONC_DATA,
                   end_row=2, end_column=_ANN_CONC_SPORE_END)

    # Riga 3: intestazioni colonne
    align_rotated = Alignment(horizontal="center", text_rotation=90)

    def _scrivi_intestazione_colonna(col, codice):
        nome = CODICI_SPECIE[codice]
        cell = ws.cell(row=3, column=col, value=nome)
        cell.font = FONT_BOLD if codice in ANNUALE_BOLD_CODICI else Font()
        cell.border = THIN_BORDER
        cell.alignment = align_rotated
        if codice in ANNUALE_VERDE_CODICI:
            cell.fill = FILL_VERDE
        elif codice in ANNUALE_VERDE_CHIARO_CODICI:
            cell.fill = FILL_VERDE_CHIARO
        else:
            cell.fill = FILL_GIALLO

    # Data (conta grezza)
    cell = ws.cell(row=3, column=1, value="Data")
    cell.font = FONT_BOLD
    cell.fill = FILL_GIALLO
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Pollini (conta grezza)
    for i, codice in enumerate(POLLINI_CODICI):
        _scrivi_intestazione_colonna(_ANN_POLL_START + i, codice)

    # Separatore
    cell = ws.cell(row=3, column=_ANN_SEP1, value="||")
    cell.font = FONT_BOLD
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Spore (conta grezza)
    for i, codice in enumerate(SPORE_CODICI):
        _scrivi_intestazione_colonna(_ANN_SPORE_START + i, codice)

    # Data (concentrazioni)
    cell = ws.cell(row=3, column=_ANN_CONC_DATA, value="Data")
    cell.font = FONT_BOLD
    cell.fill = FILL_GIALLO
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Pollini (concentrazioni)
    for i, codice in enumerate(POLLINI_CODICI):
        _scrivi_intestazione_colonna(_ANN_CONC_POLL_START + i, codice)

    # Separatore concentrazioni
    cell = ws.cell(row=3, column=_ANN_CONC_SEP, value="||")
    cell.font = FONT_BOLD
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Spore (concentrazioni)
    for i, codice in enumerate(SPORE_CODICI):
        _scrivi_intestazione_colonna(_ANN_CONC_SPORE_START + i, codice)

    # Larghezze colonne
    ws.column_dimensions["A"].width = 11
    for i in range(len(POLLINI_CODICI)):
        ws.column_dimensions[get_column_letter(_ANN_POLL_START + i)].width = 4
    ws.column_dimensions[get_column_letter(_ANN_SEP1)].width = 2
    for i in range(len(SPORE_CODICI)):
        ws.column_dimensions[get_column_letter(_ANN_SPORE_START + i)].width = 4
    ws.column_dimensions[get_column_letter(_ANN_GAP)].width = 3
    ws.column_dimensions[get_column_letter(_ANN_CONC_DATA)].width = 11
    for i in range(len(POLLINI_CODICI)):
        ws.column_dimensions[get_column_letter(_ANN_CONC_POLL_START + i)].width = 4
    ws.column_dimensions[get_column_letter(_ANN_CONC_SEP)].width = 2
    for i in range(len(SPORE_CODICI)):
        ws.column_dimensions[get_column_letter(_ANN_CONC_SPORE_START + i)].width = 4

    # Altezza riga 3 per testo ruotato
    ws.row_dimensions[3].height = 80
    # Blocca colonna Data e intestazioni
    ws.freeze_panes = "B4"
    # Auto-filter sulla riga intestazioni
    ws.auto_filter.ref = (
        f"A3:{get_column_letter(_ANN_CONC_SPORE_END)}3"
    )


def trova_riga_per_data(ws, data_str):
    """Cerca una riga nel riepilogo annuale con la data specificata.
    Ritorna numero riga o None."""
    for row in range(4, ws.max_row + 1):
        val = ws.cell(row=row, column=1).value
        if val and str(val).strip() == data_str:
            return row
    return None


def _prossima_riga_annuale(ws):
    """Trova la prima riga vuota nel riepilogo annuale (da riga 4)."""
    for row in range(4, ws.max_row + 2):
        if ws.cell(row=row, column=1).value is None:
            return row
    return ws.max_row + 1


def scrivi_riga_annuale(ws, riga, data_str, dati, fattore, modo):
    """Scrive o aggiorna una riga nel riepilogo annuale.

    modo: 'nuovo'/'sovrascrivi' = scrive i valori; 'somma' = aggiunge ai esistenti.
    """
    # Data (conta grezza)
    cell = ws.cell(row=riga, column=1, value=data_str)
    cell.border = THIN_BORDER

    # Data (concentrazioni)
    cell = ws.cell(row=riga, column=_ANN_CONC_DATA, value=data_str)
    cell.border = THIN_BORDER

    tutti_codici = POLLINI_CODICI + SPORE_CODICI
    for codice in tutti_codici:
        val_nuovo = dati.get(codice, 0)
        col_grezzo = _ann_col_grezzo(codice)
        col_conc = _ann_col_conc(codice)
        if col_grezzo is None:
            continue

        # Somma ai valori esistenti se richiesto
        if modo == "somma" and val_nuovo > 0:
            val_esistente = ws.cell(row=riga, column=col_grezzo).value
            if isinstance(val_esistente, (int, float)):
                val_nuovo = int(val_esistente) + val_nuovo

        # Conta grezza
        cell_g = ws.cell(row=riga, column=col_grezzo)
        cell_g.value = val_nuovo if val_nuovo > 0 else None
        cell_g.border = THIN_BORDER
        cell_g.alignment = ALIGN_CENTER
        if codice in ANNUALE_VERDE_CODICI:
            cell_g.fill = FILL_VERDE
        elif codice in ANNUALE_VERDE_CHIARO_CODICI:
            cell_g.fill = FILL_VERDE_CHIARO
        if codice in ANNUALE_BOLD_CODICI:
            cell_g.font = FONT_BOLD

        # Concentrazione
        conc = round(val_nuovo * fattore, 1) if val_nuovo > 0 else None
        cell_c = ws.cell(row=riga, column=col_conc)
        cell_c.value = conc
        cell_c.border = THIN_BORDER
        cell_c.alignment = ALIGN_CENTER
        cell_c.number_format = "0.0"
        if codice in ANNUALE_VERDE_CODICI:
            cell_c.fill = FILL_VERDE
        elif codice in ANNUALE_VERDE_CHIARO_CODICI:
            cell_c.fill = FILL_VERDE_CHIARO
        if codice in ANNUALE_BOLD_CODICI:
            cell_c.font = FONT_BOLD

    # Separatori
    for sep_col in (_ANN_SEP1, _ANN_CONC_SEP):
        cell = ws.cell(row=riga, column=sep_col, value="||")
        cell.border = THIN_BORDER
        cell.alignment = ALIGN_CENTER


def _nome_foglio_settimana(lunedi):
    """Ritorna il nome del foglio settimanale (es. 'W05')."""
    return f"W{lunedi.isocalendar()[1]:02d}"


def _posizione_foglio_settimana(wb_ann, nome_foglio):
    """Calcola l'indice di inserimento del foglio settimanale in ordine crescente.
    I fogli fissi (Dati Anno, Calendario) restano all'inizio."""
    nomi = wb_ann.sheetnames
    num_fissi = sum(1 for n in nomi if not re.match(r"^W\d+$", n))
    nuovo_num = int(nome_foglio[1:])
    pos = num_fissi
    for n in nomi:
        if re.match(r"^W\d+$", n) and int(n[1:]) < nuovo_num:
            pos += 1
    return pos


def crea_foglio_settimana_annuale(wb_ann, ws_riepilogo, lunedi, fattore):
    """Crea o sovrascrive il foglio settimanale nel riepilogo annuale.

    Il foglio e' nominato con il numero ISO della settimana (es. 'W05').
    Layout:
      Sezione CONTA GREZZA:  Specie x 7 giorni + Totale settimana
      Sezione CONCENTRAZIONI: stessa struttura con valori * fattore + Media
    Pollini e spore separati da una riga separatore in entrambe le sezioni.
    """
    nome_foglio = _nome_foglio_settimana(lunedi)
    domenica = lunedi + timedelta(days=6)
    settimana_num = lunedi.isocalendar()[1]

    # Crea o ricrea il foglio nella posizione ordinata
    if nome_foglio in wb_ann.sheetnames:
        idx = wb_ann.sheetnames.index(nome_foglio)
        del wb_ann[nome_foglio]
        ws_s = wb_ann.create_sheet(nome_foglio, idx)
    else:
        pos = _posizione_foglio_settimana(wb_ann, nome_foglio)
        ws_s = wb_ann.create_sheet(nome_foglio, pos)

    # ── Riga 1: Titolo ──
    titolo = (f"SETTIMANA {settimana_num}  -  "
              f"{lunedi.strftime('%d/%m/%Y')} / {domenica.strftime('%d/%m/%Y')}")
    c = ws_s.cell(row=1, column=1, value=titolo)
    c.font = FONT_BOLD_BIG
    ws_s.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)

    # Intestazioni giorni (riutilizzate in entrambe le sezioni)
    giorni_hdr = []
    for g in range(7):
        dt = lunedi + timedelta(days=g)
        giorni_hdr.append(f"{GIORNI_NOMI[g+1][:3].upper()}\n{dt.strftime('%d/%m')}")

    def _scrivi_header(riga, label_ultima, fill_hdr, font_hdr):
        """Scrive la riga intestazione colonne (specie + 7 giorni + label_ultima)."""
        c = ws_s.cell(row=riga, column=1, value="Specie")
        c.font = font_hdr; c.fill = fill_hdr; c.border = THIN_BORDER; c.alignment = ALIGN_CENTER
        for g, hdr in enumerate(giorni_hdr):
            c = ws_s.cell(row=riga, column=2 + g, value=hdr)
            c.font = font_hdr; c.fill = fill_hdr; c.border = THIN_BORDER
            c.alignment = ALIGN_CENTER_WRAP
        c = ws_s.cell(row=riga, column=9, value=label_ultima)
        c.font = font_hdr; c.fill = fill_hdr; c.border = THIN_BORDER; c.alignment = ALIGN_CENTER

    def _scrivi_specie(riga, codice, vals, valore_finale, fmt="0"):
        """Scrive una riga specie con i valori (grezzi o concentrazioni)."""
        specie = CODICI_SPECIE[codice]
        c = ws_s.cell(row=riga, column=1, value=specie)
        c.border = THIN_BORDER
        if codice in ANNUALE_VERDE_CODICI:
            c.fill = FILL_VERDE; c.font = FONT_BOLD
        elif codice in ANNUALE_VERDE_CHIARO_CODICI:
            c.fill = FILL_VERDE_CHIARO; c.font = FONT_BOLD
        for g, v in enumerate(vals):
            c = ws_s.cell(row=riga, column=2 + g,
                          value=(v if v else None))
            c.border = THIN_BORDER; c.alignment = ALIGN_CENTER; c.number_format = fmt
            if codice in ANNUALE_VERDE_CODICI:
                c.fill = FILL_VERDE
                if v: c.font = FONT_BOLD
            elif codice in ANNUALE_VERDE_CHIARO_CODICI:
                c.fill = FILL_VERDE_CHIARO
                if v: c.font = FONT_BOLD
        c = ws_s.cell(row=riga, column=9,
                      value=(valore_finale if valore_finale else None))
        c.border = THIN_BORDER; c.alignment = ALIGN_CENTER; c.number_format = fmt

    def _scrivi_separatore(riga):
        for col in range(1, 10):
            ws_s.cell(row=riga, column=col).border = THIN_BORDER

    def _scrivi_doy(riga, fill_hdr, font_hdr):
        """Scrive la riga con i numeri di giorno dell'anno (1-366)."""
        c = ws_s.cell(row=riga, column=1, value="G. anno")
        c.border = THIN_BORDER; c.fill = fill_hdr; c.font = font_hdr; c.alignment = ALIGN_CENTER
        for g in range(7):
            dt_g = lunedi + timedelta(days=g)
            doy = dt_g.timetuple().tm_yday
            c = ws_s.cell(row=riga, column=2 + g, value=doy)
            c.border = THIN_BORDER; c.fill = fill_hdr; c.font = font_hdr; c.alignment = ALIGN_CENTER
        c = ws_s.cell(row=riga, column=9)
        c.border = THIN_BORDER; c.fill = fill_hdr

    # ================================================================
    # SEZIONE A — CONTA GREZZA
    # ================================================================
    r = 2
    c = ws_s.cell(row=r, column=1, value="CONTA GREZZA")
    c.font = FONT_BOLD; c.fill = FILL_GIALLO; c.alignment = ALIGN_CENTER
    ws_s.merge_cells(start_row=r, start_column=1, end_row=r, end_column=9)

    r = 3
    _scrivi_header(r, "Totale sett.", FILL_GIALLO, FONT_BOLD)

    r = 4
    _scrivi_doy(r, FILL_GIALLO, FONT_BOLD)

    r = 5
    for codice in POLLINI_CODICI:
        row_riep = codice_to_row(codice)
        vals = [leggi_valore(ws_riepilogo, row_riep, giorno_to_col(g))
                for g in range(1, 8)]
        _scrivi_specie(r, codice, vals, sum(vals))
        r += 1

    _scrivi_separatore(r); r += 1

    for codice in SPORE_CODICI:
        row_riep = codice_to_row(codice)
        vals = [leggi_valore(ws_riepilogo, row_riep, giorno_to_col(g))
                for g in range(1, 8)]
        _scrivi_specie(r, codice, vals, sum(vals))
        r += 1

    # ================================================================
    # SEZIONE B — CONCENTRAZIONI
    # ================================================================
    r += 1  # riga vuota di separazione

    c = ws_s.cell(row=r, column=1, value="CONCENTRAZIONI (p/m3)")
    c.font = FONT_BIANCO_BOLD; c.fill = FILL_BLU; c.alignment = ALIGN_CENTER
    ws_s.merge_cells(start_row=r, start_column=1, end_row=r, end_column=9)

    r += 1
    _scrivi_header(r, "Media sett.", FILL_BLU, FONT_BIANCO_BOLD)

    r += 1
    _scrivi_doy(r, FILL_BLU, FONT_BIANCO_BOLD)

    r += 1
    for codice in POLLINI_CODICI:
        row_riep = codice_to_row(codice)
        vals_raw = [leggi_valore(ws_riepilogo, row_riep, giorno_to_col(g))
                    for g in range(1, 8)]
        conc = [round(v * fattore, 1) if v > 0 else 0.0 for v in vals_raw]
        media = round(sum(conc) / 7.0, 1)
        _scrivi_specie(r, codice, conc, media, fmt="0.0")
        r += 1

    _scrivi_separatore(r); r += 1

    for codice in SPORE_CODICI:
        row_riep = codice_to_row(codice)
        vals_raw = [leggi_valore(ws_riepilogo, row_riep, giorno_to_col(g))
                    for g in range(1, 8)]
        conc = [round(v * fattore, 1) if v > 0 else 0.0 for v in vals_raw]
        media = round(sum(conc) / 7.0, 1)
        _scrivi_specie(r, codice, conc, media, fmt="0.0")
        r += 1

    # ── Dimensioni colonne ──
    ws_s.column_dimensions["A"].width = 28
    for g in range(7):
        ws_s.column_dimensions[get_column_letter(2 + g)].width = 8
    ws_s.column_dimensions[get_column_letter(9)].width = 9
    ws_s.row_dimensions[3].height = 32

    return ws_s


def esporta_riepilogo_annuale(ws_riepilogo, lunedi, cartella):
    """Esporta i dati settimanali nel file riepilogo annuale."""
    fattore = leggi_fattore(ws_riepilogo)

    anno = lunedi.year
    nome_file = f"Riepilogo_Annuale_{anno}.xlsx"
    percorso = cartella / nome_file

    dati_settimana = raccogli_dati_giornalieri(ws_riepilogo)
    if not dati_settimana:
        print("  Nessun dato da esportare nel riepilogo annuale.")
        return

    if percorso.exists():
        try:
            wb_ann = openpyxl.load_workbook(percorso)
            ws = wb_ann.active
        except Exception as e:
            print(f"  ERRORE apertura {nome_file}: {e}")
            return
    else:
        wb_ann = openpyxl.Workbook()
        ws = wb_ann.active
        ws.title = f"Dati {anno}"
        crea_intestazione_annuale(ws, anno)

    # Foglio Calendario (crea se non esiste)
    nome_cal = "Calendario"
    if nome_cal in wb_ann.sheetnames:
        ws_cal = wb_ann[nome_cal]
    else:
        ws_cal = wb_ann.create_sheet(nome_cal)
        crea_intestazione_calendario(ws_cal, anno)

    scelta_duplicati = None
    giorni_scritti = 0

    for giorno_num in sorted(dati_settimana.keys()):
        dt = lunedi + timedelta(days=giorno_num - 1)
        data_str = dt.strftime("%d/%m/%Y")
        dati = dati_settimana[giorno_num]

        # ── Foglio Dati (righe = giorni) ──
        riga_esistente = trova_riga_per_data(ws, data_str)

        if riga_esistente:
            if scelta_duplicati is None:
                print(f"\n  Il giorno '{data_str}' e' gia' presente nel riepilogo.")
                print("  Come gestire i duplicati?")
                print("    a) Sovrascrivere i dati esistenti")
                print("    b) Aggiungere una nuova riga")
                print("    c) Sommare ai dati esistenti")
                while True:
                    risp = input("  Scelta (a/b/c): ").strip().lower()
                    if risp in ("a", "b", "c"):
                        scelta_duplicati = risp
                        break
                    print("  Scelta non valida.")

            if scelta_duplicati == "a":
                scrivi_riga_annuale(ws, riga_esistente, data_str, dati,
                                    fattore, "sovrascrivi")
            elif scelta_duplicati == "b":
                scrivi_riga_annuale(ws, _prossima_riga_annuale(ws), data_str,
                                    dati, fattore, "nuovo")
            else:
                scrivi_riga_annuale(ws, riga_esistente, data_str, dati,
                                    fattore, "somma")
        else:
            scrivi_riga_annuale(ws, _prossima_riga_annuale(ws), data_str,
                                dati, fattore, "nuovo")

        # ── Foglio Calendario (colonne = giorni) ──
        col_esistente = trova_colonna_per_data_calendario(ws_cal, data_str)

        if col_esistente:
            modo_cal = {"a": "sovrascrivi", "b": "nuovo", "c": "somma"}.get(
                scelta_duplicati, "sovrascrivi"
            )
            if scelta_duplicati == "b":
                scrivi_colonna_calendario(ws_cal, _prossima_colonna_calendario(ws_cal),
                                          data_str, dati, fattore, "nuovo")
            else:
                scrivi_colonna_calendario(ws_cal, col_esistente,
                                          data_str, dati, fattore, modo_cal)
        else:
            scrivi_colonna_calendario(ws_cal, _prossima_colonna_calendario(ws_cal),
                                      data_str, dati, fattore, "nuovo")

        giorni_scritti += 1

    # ── Foglio settimanale (W##) ──
    crea_foglio_settimana_annuale(wb_ann, ws_riepilogo, lunedi, fattore)

    try:
        wb_ann.save(percorso)
        wb_ann.close()
        nome_sett = _nome_foglio_settimana(lunedi)
        print(f"\n  [OK] Riepilogo annuale aggiornato: {percorso}")
        print(f"       {giorni_scritti} giorni esportati, fattore={fattore}")
        print(f"       Foglio settimanale: {nome_sett}")
    except Exception as e:
        print(f"  ERRORE salvataggio {nome_file}: {e}")


def genera_bollettini_word(ws_riepilogo, lunedi, lunedi_str, cartella):
    """Genera i bollettini Word (ITA e ENG) a partire dai dati del foglio riepilogo."""
    if _docx_module is None:
        print("  ATTENZIONE: python3-docx non installato. Bollettini Word non generati.")
        print("  Installa con: sudo apt install python3-docx")
        return

    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor, Pt

    # Fattore di conversione
    fattore = leggi_fattore(ws_riepilogo)

    # Carica soglie dinamiche (dal workbook o dal file esterno)
    soglie = carica_soglie(ws_riepilogo.parent) or carica_soglie() or {}

    # Calcola concentrazioni per ogni riga e filtra quelle con almeno un valore > 0
    righe_dati = []
    for ita_nome, eng_nome, riep_rows, soglie_fam, fb_ass, fb_bas, fb_med in BOLL_WORD_RIGHE:
        assente_max, bassa_max, media_max = soglie.get(soglie_fam, (fb_ass, fb_bas, fb_med))
        conc_giorni = []
        for g in range(1, 8):
            col = giorno_to_col(g)
            raw = sum(leggi_valore(ws_riepilogo, r, col) for r in riep_rows)
            conc_giorni.append(round(raw * fattore, 1))
        if any(v > 0 for v in conc_giorni):
            media = round(sum(conc_giorni) / 7, 1)
            righe_dati.append((ita_nome, eng_nome, conc_giorni, media,
                                assente_max, bassa_max, media_max))

    if not righe_dati:
        print("  Nessun dato presente: bollettini Word non generati.")
        return

    def _colore(valore, assente_max, bassa_max, media_max):
        if valore > media_max:   return "C00000"
        if valore > bassa_max:   return "FFC000"
        if valore > assente_max: return "FFFF00"
        return "92D050"

    # Larghezze colonne in dxa: [specie, lun, mar, mer, gio, ven, sab, dom, media, tendenza]
    # Totale 15398 dxa ≈ 27.16 cm (area di testo pagina landscape con margini 1.27 cm)
    _COL_WIDTHS = [2000, 1300, 1300, 1300, 1300, 1300, 1300, 1300, 1800, 2498]

    def _set_table_widths(table):
        """Aggiorna tblGrid e tblW per ridistribuire le larghezze colonne."""
        tbl = table._tbl
        tblGrid = tbl.find(qn("w:tblGrid"))
        if tblGrid is not None:
            for i, gc in enumerate(tblGrid.findall(qn("w:gridCol"))):
                if i < len(_COL_WIDTHS):
                    gc.set(qn("w:w"), str(_COL_WIDTHS[i]))
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblW = tblPr.find(qn("w:tblW"))
            if tblW is not None:
                tblW.set(qn("w:w"), str(sum(_COL_WIDTHS)))
                tblW.set(qn("w:type"), "dxa")

    def _set_cell_color(cell, rgb_hex):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), rgb_hex)

    def _set_cell_borders(cell):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        for side in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:color"), "000000")
            border.set(qn("w:sz"), "4")
            tcBorders.append(border)

    def _set_cell_paragraph_format(cell):
        """Imposta allineamento verticale centrato e allineamento paragrafo centrato."""
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        vAlign = tcPr.find(qn("w:vAlign"))
        if vAlign is None:
            vAlign = OxmlElement("w:vAlign")
            tcPr.append(vAlign)
        vAlign.set(qn("w:val"), "center")
        para = cell.paragraphs[0]
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            para._p.insert(0, pPr)
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

    def _set_row_height(row, height_twips):
        """Imposta altezza fissa riga in twips."""
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            row._tr.insert(0, trPr)
        trH = trPr.find(qn("w:trHeight"))
        if trH is None:
            trH = OxmlElement("w:trHeight")
            trPr.append(trH)
        trH.set(qn("w:val"), str(height_twips))
        trH.set(qn("w:hRule"), "exact")

    def _set_cell_text(cell, text, italic=False, color=None, bold=None, size_pt=None):
        para = cell.paragraphs[0]
        for r in para._p.findall(qn("w:r")):
            para._p.remove(r)
        run = para.add_run(text)
        run.font.name = "Arial Narrow"
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
        if bold is not None:
            run.font.bold = bold
        if size_pt is not None:
            run.font.size = Pt(size_pt)

    for lang in ("ITA", "ENG"):
        template_name = f"{lang}_Template_Bollettino_pubblicazione.docx"
        template_path = SCRIPT_DIR / template_name
        if not template_path.exists():
            print(f"  ATTENZIONE: template non trovato: {template_path}")
            continue

        doc = Document(str(template_path))
        table = doc.tables[0]
        tbl = table._tbl
        _set_table_widths(table)

        # Aggiorna intestazione (riga 0)
        header_cells = table.rows[0].cells
        if lang == "ITA":
            _set_cell_text(header_cells[0], f"POLLINI \u2013 {_MESI_ITA[lunedi.month - 1]} {lunedi.year}",
                           color="002060", bold=True, size_pt=10)
            giorni_long = _GIORNI_ITA_LONG
        else:
            _set_cell_text(header_cells[0], f"POLLEN \u2013 {lunedi.strftime('%B')} {lunedi.year}",
                           color="002060", bold=True, size_pt=10)
            giorni_long = _GIORNI_ENG_LONG

        for i in range(7):
            giorno_dt = lunedi + timedelta(days=i)
            _set_cell_text(header_cells[i + 1], f"{giorni_long[i]} {giorno_dt.day}",
                           color="002060", bold=True, size_pt=9)

        # Rimuovi tutte le righe dati esistenti (mantieni solo l'intestazione)
        for row in list(table.rows[1:]):
            tbl.remove(row._tr)

        # Aggiungi righe con i dati
        for ita_nome, eng_nome, conc_giorni, media, assente_max, bassa_max, media_max in righe_dati:
            nome = ita_nome if lang == "ITA" else eng_nome
            new_row = table.add_row()
            cells = new_row.cells

            # Formattazione base su tutte le celle
            for cell in cells:
                _set_cell_borders(cell)
                _set_cell_paragraph_format(cell)

            # Colonna 0: nome specie (corsivo, blu, 9pt)
            _set_cell_text(cells[0], nome, italic=True, color="002060", size_pt=9)

            # Colonne 1-7: colore per ogni giorno (nessun testo)
            for i, conc in enumerate(conc_giorni):
                _set_cell_color(cells[i + 1], _colore(conc, assente_max, bassa_max, media_max))

            # Colonna 8: colore media
            _set_cell_color(cells[8], _colore(media, assente_max, bassa_max, media_max))

            # Colonna 9: tendenza (vuota, da compilare manualmente)

            # Altezza riga fissa (360 twips = 18pt, come nel template)
            _set_row_height(new_row, 360)

        # Salva
        output_name = f"Bollettino_{lang}_{lunedi_str}.docx"
        output_path = cartella / output_name
        doc.save(str(output_path))
        print(f"  Bollettino {lang}: {output_path}")

    print("  Bollettini Word generati.")


# ============================================================
# Foglio Calendario (trasposto: specie in righe, date in colonne)
# ============================================================
def _cal_row_for_codice(codice):
    """Riga nel foglio Calendario per un codice specie."""
    n = int(codice)
    if 1 <= n <= 47:
        return _CAL_DATA_START_ROW + (n - 1)   # righe 4-50
    if 48 <= n <= 59:
        return _CAL_SEP_ROW + 1 + (n - 48)     # righe 52-63
    return None


def crea_intestazione_calendario(ws, anno):
    """Crea intestazione del foglio Calendario (righe 1-3 + specie in colonna A)."""
    # Riga 1: titolo
    cell = ws.cell(row=1, column=1, value=f"CALENDARIO POLLINICO {anno}")
    cell.font = FONT_BOLD_BIG
    cell.alignment = ALIGN_CENTER

    # Riga 2: sezione
    cell = ws.cell(row=2, column=1, value="CONCENTRAZIONI (p/m3)")
    cell.font = FONT_BOLD
    cell.alignment = ALIGN_CENTER
    cell.fill = FILL_GIALLO

    # Riga 3 col A: intestazione date
    cell = ws.cell(row=_CAL_HEADER_ROW, column=1, value="Specie")
    cell.font = FONT_BOLD
    cell.fill = FILL_GIALLO
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Riga 4 col A: etichetta giorno dell'anno
    cell = ws.cell(row=_CAL_DOY_ROW, column=1, value="G. anno")
    cell.font = FONT_BOLD
    cell.fill = FILL_GIALLO
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Righe pollini
    for i, codice in enumerate(POLLINI_CODICI):
        row = _CAL_DATA_START_ROW + i
        cell = ws.cell(row=row, column=1, value=CODICI_SPECIE[codice])
        cell.font = FONT_BOLD if codice in ANNUALE_BOLD_CODICI else Font()
        cell.border = THIN_BORDER
        if codice in ANNUALE_VERDE_CODICI:
            cell.fill = FILL_VERDE
        elif codice in ANNUALE_VERDE_CHIARO_CODICI:
            cell.fill = FILL_VERDE_CHIARO

    # Riga separatore
    cell = ws.cell(row=_CAL_SEP_ROW, column=1, value="||")
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Righe spore
    for i, codice in enumerate(SPORE_CODICI):
        row = _CAL_SEP_ROW + 1 + i
        cell = ws.cell(row=row, column=1, value=CODICI_SPECIE[codice])
        cell.font = FONT_BOLD if codice in ANNUALE_BOLD_CODICI else Font()
        cell.border = THIN_BORDER
        if codice in ANNUALE_VERDE_CHIARO_CODICI:
            cell.fill = FILL_VERDE_CHIARO

    # Larghezza colonna specie, altezze righe, freeze
    ws.column_dimensions["A"].width = 26
    ws.row_dimensions[_CAL_HEADER_ROW].height = 60
    ws.row_dimensions[_CAL_DOY_ROW].height = 16
    ws.freeze_panes = "B5"


def trova_colonna_per_data_calendario(ws, data_str):
    """Cerca la colonna della data nel foglio Calendario. Ritorna colonna o None."""
    for col in range(_CAL_COL_DATE_START, ws.max_column + 1):
        val = ws.cell(row=_CAL_HEADER_ROW, column=col).value
        if val and str(val).strip() == data_str:
            return col
    return None


def _prossima_colonna_calendario(ws):
    """Prima colonna vuota nel foglio Calendario (da col 2)."""
    for col in range(_CAL_COL_DATE_START, ws.max_column + 2):
        if ws.cell(row=_CAL_HEADER_ROW, column=col).value is None:
            return col
    return ws.max_column + 1


def scrivi_colonna_calendario(ws, col, data_str, dati, fattore, modo):
    """Scrive o aggiorna una colonna nel foglio Calendario."""
    align_rotated = Alignment(horizontal="center", text_rotation=90)

    # Intestazione data (ruotata)
    cell = ws.cell(row=_CAL_HEADER_ROW, column=col, value=data_str)
    cell.font = FONT_BOLD
    cell.fill = FILL_GIALLO
    cell.border = THIN_BORDER
    cell.alignment = align_rotated
    ws.column_dimensions[get_column_letter(col)].width = 5

    # Giorno dell'anno
    try:
        dt_col = datetime.strptime(data_str, "%d/%m/%Y")
        doy = dt_col.timetuple().tm_yday
        cell_doy = ws.cell(row=_CAL_DOY_ROW, column=col, value=doy)
        cell_doy.font = FONT_BOLD
        cell_doy.fill = FILL_GIALLO
        cell_doy.border = THIN_BORDER
        cell_doy.alignment = ALIGN_CENTER
    except ValueError:
        pass

    # Separatore
    cell = ws.cell(row=_CAL_SEP_ROW, column=col, value="||")
    cell.border = THIN_BORDER
    cell.alignment = ALIGN_CENTER

    # Concentrazioni specie
    tutti_codici = POLLINI_CODICI + SPORE_CODICI
    for codice in tutti_codici:
        val_nuovo = dati.get(codice, 0)
        row = _cal_row_for_codice(codice)
        if row is None:
            continue

        if modo == "somma" and val_nuovo > 0:
            conc_esistente = ws.cell(row=row, column=col).value
            if isinstance(conc_esistente, (int, float)):
                conc = round(conc_esistente + val_nuovo * fattore, 1)
            else:
                conc = round(val_nuovo * fattore, 1)
        else:
            conc = round(val_nuovo * fattore, 1) if val_nuovo > 0 else None

        cell = ws.cell(row=row, column=col)
        cell.value = conc
        cell.border = THIN_BORDER
        cell.alignment = ALIGN_CENTER
        cell.number_format = "0.0"
        if codice in ANNUALE_VERDE_CODICI:
            cell.fill = FILL_VERDE
        elif codice in ANNUALE_VERDE_CHIARO_CODICI:
            cell.fill = FILL_VERDE_CHIARO
        if codice in ANNUALE_BOLD_CODICI and conc:
            cell.font = FONT_BOLD


# ============================================================
# Riepiloghi e correzioni
# ============================================================
def controlla_giorno_esistente(ws, giorno_num):
    """Controlla se la colonna del giorno ha gia' dati. Ritorna il totale."""
    col = giorno_to_col(giorno_num)
    return sum(
        leggi_valore(ws, r, col)
        for r in list(range(6, 53)) + list(range(58, 70))
    )


def mostra_riepilogo_giorno(ws, giorno_num):
    """Mostra riepilogo delle specie con conteggio > 0 per il giorno,
    separando pollini e spore."""
    col = giorno_to_col(giorno_num)
    nome_giorno = GIORNI_NOMI[giorno_num].upper()
    print(f"\n  Riepilogo {nome_giorno}:")

    totale_pollini = 0
    totale_spore = 0
    has_pollini = False
    has_spore = False

    # Pollini (01-47)
    print("    POLLINI:")
    for codice_str, specie in CODICI_SPECIE.items():
        n = int(codice_str)
        if n > 47:
            continue
        row = codice_to_row(codice_str)
        if row is None:
            continue
        val = leggi_valore(ws, row, col)
        if val > 0:
            print(f"      [{codice_str}] {specie}: {val}")
            totale_pollini += val
            has_pollini = True
    if not has_pollini:
        print("      (nessun dato)")
    print(f"      --- Totale pollini: {totale_pollini}")

    # Spore (48-59)
    print("    SPORE:")
    for codice_str, specie in CODICI_SPECIE.items():
        n = int(codice_str)
        if n < 48:
            continue
        row = codice_to_row(codice_str)
        if row is None:
            continue
        val = leggi_valore(ws, row, col)
        if val > 0:
            print(f"      [{codice_str}] {specie}: {val}")
            totale_spore += val
            has_spore = True
    if not has_spore:
        print("      (nessun dato)")
    print(f"      --- Totale spore: {totale_spore}")

    totale = totale_pollini + totale_spore
    if totale > 0:
        print(f"    === TOTALE GIORNO: {totale}")
    print()


def mostra_riepilogo_settimana(ws):
    """Mostra tabella riassuntiva lun-dom con tutte le specie che hanno dati."""
    intestazione = "  " + " " * 28 + "LUN  MAR  MER  GIO  VEN  SAB  DOM"
    print(f"\n{intestazione}")
    print("  " + "-" * 63)

    def _stampa_sezione(codici_range, etichetta_totale):
        totali = [0] * 7
        has_data = False
        for codice_str, specie in CODICI_SPECIE.items():
            n = int(codice_str)
            if n not in codici_range:
                continue
            row = codice_to_row(codice_str)
            vals = [leggi_valore(ws, row, giorno_to_col(g)) for g in range(1, 8)]
            if any(v > 0 for v in vals):
                has_data = True
                sp = specie[:22]
                vs = "".join(f"{v:5}" if v > 0 else "    -" for v in vals)
                print(f"  [{codice_str}] {sp:<24} {vs}")
                for i in range(7):
                    totali[i] += vals[i]
        if not has_data:
            print(f"  (nessun dato)")
        ts = "".join(f"{v:5}" for v in totali)
        print(f"  {etichetta_totale:<28} {ts}")

    _stampa_sezione(range(1, 48), "--- Totale pollini")
    print()
    _stampa_sezione(range(48, 60), "--- Totale spore")
    print()


def mostra_storico(storico):
    """Mostra gli ultimi 10 inserimenti."""
    if not storico:
        print("\n  (nessun inserimento in questa sessione)\n")
        return
    print("\n  Ultimi inserimenti:")
    for data, codice, specie, quantita, ora in storico[-10:]:
        if quantita > 1:
            print(f"    {ora}  [{codice}] {specie} x{quantita}  ({data})")
        else:
            print(f"    {ora}  [{codice}] {specie}  ({data})")
    print()


def aggiungi_nota(ws_log, log_row, data_str):
    """Permette di aggiungere una nota testuale per la giornata."""
    nota = input("  Nota: ").strip()
    if not nota:
        print("  (nessuna nota inserita)")
        return log_row
    scrivi_log(ws_log, log_row, data_str, "--", "NOTA", nota=nota)
    log_row += 1
    print(f"  Nota registrata: {nota}")
    return log_row


def correggi_giorno(ws_riepilogo, ws_log, lunedi, log_row):
    """Permette di correggere i dati di un giorno qualsiasi."""
    print("\nQuale giorno vuoi correggere?")
    for num, nome in GIORNI_NOMI.items():
        data_giorno = lunedi + timedelta(days=num - 1)
        print(f"  {num}) {nome.upper():<12} {data_giorno.strftime('%d-%m-%Y')}")

    while True:
        scelta = input("Scegli (1-7, invio per annullare): ").strip()
        if not scelta:
            return log_row
        if scelta in [str(x) for x in range(1, 8)]:
            giorno_num = int(scelta)
            break
        print("Scelta non valida, riprova.")

    col = giorno_to_col(giorno_num)
    nome_giorno = GIORNI_NOMI[giorno_num].upper()
    data_str = (lunedi + timedelta(days=giorno_num - 1)).strftime("%d-%m-%Y")

    # Mostra dati attuali
    print(f"\n  Dati {nome_giorno} {data_str}:")
    found = False
    for codice_str, specie in CODICI_SPECIE.items():
        row = codice_to_row(codice_str)
        if row is None:
            continue
        val = leggi_valore(ws_riepilogo, row, col)
        if val > 0:
            print(f"    [{codice_str}] {specie}: {val}")
            found = True

    if not found:
        print("    (nessun dato)")
        return log_row

    codice = normalizza_codice(input("\n  Codice da correggere (invio per annullare): ").strip())
    if not codice:
        return log_row
    if codice not in CODICI_SPECIE:
        print(f"  Codice non riconosciuto: {codice}")
        return log_row

    specie = CODICI_SPECIE[codice]
    row = codice_to_row(codice)
    val_attuale = leggi_valore(ws_riepilogo, row, col)

    nuovo = input(f"  {specie}: {val_attuale} -> nuovo valore: ").strip()
    if not nuovo:
        return log_row
    try:
        nuovo_val = int(nuovo)
    except ValueError:
        print("  Valore non valido.")
        return log_row
    if nuovo_val < 0:
        print("  Il valore non puo' essere negativo.")
        return log_row

    ws_riepilogo.cell(row=row, column=col, value=nuovo_val)
    scrivi_log(ws_log, log_row, data_str, codice, specie,
               nota=f"CORREZIONE {val_attuale}->{nuovo_val}")
    log_row += 1

    print(f"  Corretto: [{codice}] {specie}: {val_attuale} -> {nuovo_val}")
    return log_row


# ============================================================
# Autosave e pulizia
# ============================================================
_autosave_thread = None
_autosave_lock = threading.Lock()


def autosave(wb, lunedi_str, sincrono=False):
    """Salva su file autosave. Se sincrono=False (default) usa un thread background.
    Il lock serializza il salvataggio rispetto alle modifiche del main thread.
    Se un autosave e' gia' in corso, salta (evita lag da save sovrapposti)."""
    global _autosave_thread
    if not _autosave_lock.acquire(blocking=False):
        return
    path = OUTPUT_DIR / f"~autosave_{lunedi_str}.xlsx"

    def _do_save():
        try:
            # Salva su file temporaneo, poi rinomina atomicamente.
            # Evita che una lettura concorrente (GUI o prossimo avvio) trovi
            # un archivio ZIP troncato ("truncated header" / "not a zip file").
            tmp = path.with_suffix(".tmp")
            wb.save(tmp)
            tmp.replace(path)
        except Exception:
            try:
                tmp.unlink(missing_ok=True)
            except Exception:
                pass
        finally:
            _autosave_lock.release()

    if sincrono:
        _do_save()
    else:
        _autosave_thread = threading.Thread(target=_do_save, daemon=True)
        _autosave_thread.start()


def _attendi_autosave():
    """Attende il completamento dell'eventuale autosave in corso (max 10s)."""
    if _autosave_thread and _autosave_thread.is_alive():
        _autosave_thread.join(timeout=10)


def _safe_remove(path, tentativi=3, pausa=0.5):
    """Elimina un file con retry per gestire lock di Windows (sharing violation)."""
    for i in range(tentativi):
        try:
            path.unlink()
            return True
        except PermissionError:
            if i < tentativi - 1:
                time.sleep(pausa)
    print(f"  Attenzione: impossibile eliminare {path.name} (file in uso).")
    return False


def _safe_rename(src, dst, tentativi=3, pausa=0.5):
    """Rinomina un file con retry per gestire lock di Windows."""
    for i in range(tentativi):
        try:
            if dst.exists():
                dst.unlink()
            src.rename(dst)
            return True
        except PermissionError:
            if i < tentativi - 1:
                time.sleep(pausa)
    print(f"  Attenzione: impossibile rinominare {src.name} (file in uso).")
    return False


def pulisci_file_temporanei(lunedi_str, file_ripreso=None, salvataggio_ok=False):
    """Gestisce i file temporanei dopo la chiusura dello script.

    Se salvataggio_ok=True: chiede se eliminare l'autosave; rimuove il file ripreso.
    Se salvataggio_ok=False: chiede se conservare l'autosave per ripresa futura.
    """
    autosave_path = OUTPUT_DIR / f"~autosave_{lunedi_str}.xlsx"

    if salvataggio_ok:
        # Salvataggio riuscito: elimina sempre l'autosave (il file definitivo e' gia' stato salvato)
        if autosave_path.exists():
            _safe_remove(autosave_path)
        # Cancella sempre il file temporaneo ripreso (gia' sostituito dal salvataggio)
        if file_ripreso and file_ripreso.name.startswith(("~autosave_", "incompleto_")):
            if file_ripreso.exists():
                _safe_remove(file_ripreso)
    else:
        # Uscita senza salvare: chiedi se conservare il lavoro per ripresa
        if autosave_path.exists():
            risp = input("\n  Conservare il lavoro per riprenderlo alla prossima sessione? (s/n): ").strip().lower()
            if risp == "s":
                incompleto_path = OUTPUT_DIR / f"incompleto_{lunedi_str}.xlsx"
                if _safe_rename(autosave_path, incompleto_path):
                    print(f"  Il lavoro e' salvato in: {incompleto_path.name}")
                    print(f"  Puoi riprenderlo alla prossima esecuzione.")
                else:
                    print(f"  Il lavoro e' in: {autosave_path.name}")
                    print(f"  Puoi riprenderlo alla prossima esecuzione.")
            else:
                _safe_remove(autosave_path)
        # Non toccare file_ripreso se l'utente era partito da un incompleto


# ============================================================
# Operazioni di sessione: inserimento e undo
# ============================================================
def esegui_inserimento(ws_riepilogo, ws_log, col, data_str, codice, quantita,
                       log_row, stato):
    """Inserisce una specie nel riepilogo e nel log.
    Ritorna log_row aggiornato.
    """
    specie = CODICI_SPECIE[codice]
    row_riep = codice_to_row(codice)

    val = leggi_valore(ws_riepilogo, row_riep, col)
    ws_riepilogo.cell(row=row_riep, column=col, value=val + quantita)

    for _ in range(quantita):
        scrivi_log(ws_log, log_row, data_str, codice, specie)
        log_row += 1

    stato["ultimo_codice"] = codice
    stato["storico"].append((data_str, codice, specie, quantita,
                             datetime.now().strftime("%H:%M:%S")))

    new_val = val + quantita
    if quantita > 1:
        print(f"  -> [{codice}] {specie} x{quantita}  (totale giorno: {new_val})")
    else:
        print(f"  -> [{codice}] {specie}  (totale giorno: {new_val})")

    if stato["beep"]:
        _beep()

    return log_row


def esegui_undo(ws_riepilogo, ws_log, col, data_str, undo_stack, log_row):
    """Annulla l'ultimo inserimento.
    Ritorna (log_row, quantita_annullata).
    """
    last_codice, last_qty = undo_stack.pop()
    last_specie = CODICI_SPECIE[last_codice]
    row_riep = codice_to_row(last_codice)

    val = leggi_valore(ws_riepilogo, row_riep, col)
    if val >= last_qty:
        ws_riepilogo.cell(row=row_riep, column=col, value=val - last_qty)

    for _ in range(last_qty):
        log_row -= 1
        cancella_riga_log(ws_log, log_row)

    scrivi_log(ws_log, log_row, data_str, last_codice, last_specie,
               nota=f"ANNULLATO x{last_qty}")
    log_row += 1

    if last_qty > 1:
        print(f"  <- Annullato: [{last_codice}] {last_specie} x{last_qty}")
    else:
        print(f"  <- Annullato: [{last_codice}] {last_specie}")

    return log_row, last_qty


# ============================================================
# Sessione di inserimento per un giorno
# ============================================================
def sessione_giorno(ctx, giorno_num, data_str, log_row, stato):
    """Gestisce l'inserimento per un singolo giorno.

    ctx e' un dizionario con il contesto invariante della sessione:
        ws_riepilogo, ws_log, wb, lunedi, lunedi_str,
        prima_data, nome_ripreso, file_ripreso

    stato e' un dizionario mutabile che mantiene:
        - ultimo_codice: per il comando '.'
        - storico: lista inserimenti per il comando 'l'
        - beep: flag beep sonoro
        - file_salvato: flag che indica se il file è stato salvato durante l'uscita

    Ritorna:
        ("continue", log_row) o ("quit", log_row)
    """
    ws_riepilogo = ctx["ws_riepilogo"]
    ws_log = ctx["ws_log"]
    wb = ctx["wb"]
    lunedi = ctx["lunedi"]
    lunedi_str = ctx["lunedi_str"]
    prima_data = ctx["prima_data"]
    nome_ripreso = ctx["nome_ripreso"]
    file_ripreso = ctx["file_ripreso"]

    col = giorno_to_col(giorno_num)
    nome_giorno = GIORNI_NOMI[giorno_num].upper()
    abbrev = GIORNI_NOMI[giorno_num][:3].upper()

    print(f"\n  Giorno:   {nome_giorno}")
    print(f"  Data:     {data_str}")
    print(f"  Colonna:  {chr(ord('A') + col - 1)} (riepilogo settimanale)")
    print()

    # Autosave immediato: crea subito il file autosave cosi' la GUI puo'
    # iniziare a tracciarlo senza aspettare i primi 5 inserimenti.
    autosave(wb, lunedi_str)
    print(f"  [auto-salvato]: {OUTPUT_DIR / f'~autosave_{lunedi_str}.xlsx'}")

    conteggio = controlla_giorno_esistente(ws_riepilogo, giorno_num)
    conteggio_autosave = 0
    undo_stack = []       # lista di (codice, quantita)
    undo_consecutivi = 0

    while True:
        try:
            user_input = input(f"{abbrev} [{conteggio}] >> ").strip()
            if not user_input:
                continue
            cmd = user_input.lower()

            # ── Comandi di uscita ──
            if cmd == "q":
                print(f"\n  Chiusura {nome_giorno}: {conteggio} osservazioni.")
                salvataggio_eseguito, cartella_salvata = menu_uscita_salvataggio(wb, prima_data, nome_ripreso, file_ripreso, stato)
                stato["file_salvato"] = salvataggio_eseguito
                stato["cartella_salvata"] = cartella_salvata
                return "quit", log_row

            if cmd == "d":
                print(f"\n  Chiusura {nome_giorno}: {conteggio} osservazioni.")
                return "continue", log_row

            # ── Comandi informativi ──
            if cmd == "h":
                display_menu()
                continue
            if cmd == "r":
                mostra_riepilogo_giorno(ws_riepilogo, giorno_num)
                continue
            if cmd == "w":
                mostra_riepilogo_settimana(ws_riepilogo)
                continue
            if cmd == "l":
                mostra_storico(stato["storico"])
                continue
            if cmd == "b":
                stato["beep"] = not stato["beep"]
                status = "ATTIVO" if stato["beep"] else "disattivo"
                print(f"  Beep sonoro: {status}")
                if stato["beep"]:
                    _beep()
                continue
            if cmd == "c":
                log_row = correggi_giorno(ws_riepilogo, ws_log, lunedi, log_row)
                continue
            if cmd == "n":
                log_row = aggiungi_nota(ws_log, log_row, data_str)
                continue

            # ── Salvataggio mid-session ──
            if cmd == "s":
                _attendi_autosave()
                if stato.get("percorso_salvato"):
                    wb.save(stato["percorso_salvato"])
                    print(f"  [OK] File salvato: {stato['percorso_salvato']}")
                else:
                    s_data = re.sub(r"^(Conta_Pollinica_|~autosave_|incompleto_)+", "", prima_data)
                    s_nome = nome_ripreso or f"Conta_Pollinica_{s_data or prima_data}.xlsx"
                    s_percorso = chiedi_percorso_salvataggio(s_nome)
                    if s_percorso:
                        wb.save(s_percorso)
                        stato["percorso_salvato"] = s_percorso
                        print(f"  [OK] File salvato: {s_percorso}")
                    else:
                        print("  Salvataggio annullato.")
                continue

            # ── Undo ──
            if cmd == "u":
                if not undo_stack:
                    print("  Nessun inserimento da annullare.")
                    continue
                undo_consecutivi += 1
                if undo_consecutivi > 5 and (undo_consecutivi - 1) % 5 == 0:
                    risp = input(
                        f"  Hai annullato {undo_consecutivi - 1} inserimenti "
                        f"di fila. Continuare? (s/n): "
                    ).strip().lower()
                    if risp != "s":
                        undo_consecutivi = 0
                        continue
                codice_undo = undo_stack[-1][0]  # peek prima che esegui_undo lo rimuova
                log_row, qty = esegui_undo(
                    ws_riepilogo, ws_log, col, data_str, undo_stack, log_row
                )
                conteggio -= qty
                conteggio_autosave = max(0, conteggio_autosave - qty)
                # Marker per aggiornamento live dopo undo
                _val = leggi_valore(ws_riepilogo, codice_to_row(codice_undo), col)
                print(f"__GUI_DELTA__|{codice_undo}|{giorno_num}|{_val}", flush=True)
                continue

            # Reset contatore undo consecutivi
            undo_consecutivi = 0

            # ── Determinare codice e quantita' ──
            if cmd == ".":
                if not stato["ultimo_codice"]:
                    print("  Nessun codice precedente da ripetere.")
                    continue
                codice = stato["ultimo_codice"]
                quantita = 1
            else:
                quantita = 1
                codice = user_input
                match = re.match(r"^(\d{1,2})[xX*](\d+)$", user_input)
                if match:
                    codice = match.group(1)
                    quantita = int(match.group(2))
                    if quantita < 1 or quantita > 100:
                        print("  Quantita' non valida (1-100).")
                        continue

                codice = normalizza_codice(codice)
                if codice not in CODICI_SPECIE:
                    print(f"  Codice non riconosciuto: {user_input}")
                    continue

            # ── Inserimento ──
            log_row = esegui_inserimento(
                ws_riepilogo, ws_log, col, data_str, codice, quantita,
                log_row, stato,
            )
            undo_stack.append((codice, quantita))
            conteggio += quantita
            conteggio_autosave += quantita
            # Marker per aggiornamento live delle tabelle GUI (senza leggere il file)
            _val = leggi_valore(ws_riepilogo, codice_to_row(codice), col)
            print(f"__GUI_DELTA__|{codice}|{giorno_num}|{_val}", flush=True)

            if conteggio_autosave % AUTOSAVE_INTERVAL == 0:
                autosave(wb, lunedi_str)
                print(f"  [auto-salvato]: {OUTPUT_DIR / f'~autosave_{lunedi_str}.xlsx'}")

        except KeyboardInterrupt:
            _attendi_autosave()
            autosave(wb, lunedi_str, sincrono=True)
            print(f"\n\n  Interrotto. {nome_giorno}: {conteggio} osservazioni.")
            print(f"  Auto-salvato su ~autosave_{lunedi_str}.xlsx")
            return "quit", log_row

    return "quit", log_row


# ============================================================
# Configurazione cartella di lavoro
# ============================================================
def carica_o_crea_config():
    """Legge la cartella di lavoro per l'anno corrente dal file pollencounter.cfg.

    Se il file non esiste o non ha una voce per quest'anno, chiede la cartella
    all'utente e la salva. Ritorna un Path valido.
    """
    anno = datetime.now().year
    chiave = str(anno)
    config = {}

    if CONFIG_FILE.exists():
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                config = json.load(f)
        except Exception:
            config = {}

    if chiave in config:
        cartella = Path(config[chiave])
        if cartella.exists():
            print(f"  Cartella di lavoro {anno}: {cartella}")
            return cartella
        print(f"  Attenzione: la cartella configurata non esiste piu': {cartella}")

    # Nessuna configurazione valida per quest'anno: chiedere all'utente
    print(f"\nConfigurazione cartella di lavoro per l'anno {anno}.")
    print(f"  Tutti i file di questa stagione (settimanali e riepilogo annuale)")
    print(f"  verranno salvati in questa cartella.")
    print(f"  Invio = cartella corrente ({OUTPUT_DIR})")
    print(f"__GUI_ASKDIR__|{OUTPUT_DIR}", flush=True)
    risposta = input(
        "  Inserisci il percorso (o premi Invio per usare la cartella corrente): "
    ).strip()

    if not risposta:
        cartella = OUTPUT_DIR
    else:
        cartella = Path(risposta)
        try:
            cartella.mkdir(parents=True, exist_ok=True)
        except Exception:
            print("  Percorso non valido, uso la cartella corrente.")
            cartella = OUTPUT_DIR

    config[chiave] = str(cartella)
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        print(f"  [OK] Configurazione salvata in: {CONFIG_FILE.name}")
    except Exception as e:
        print(f"  Attenzione: impossibile salvare la configurazione: {e}")

    return cartella


# ============================================================
# Main
# ============================================================
def main():
    global OUTPUT_DIR

    if not TEMPLATE_FILE.exists():
        print(f"ERRORE: Template non trovato: {TEMPLATE_FILE}")
        sys.exit(1)

    print("=" * 60)
    print("SISTEMA CONTA POLLINICA - RIEPILOGO SETTIMANALE")
    print("=" * 60)

    # Carica (o chiede) la cartella di lavoro per l'anno corrente
    OUTPUT_DIR = carica_o_crea_config()

    # Cerca file esistenti per eventuale ripresa
    file_esistenti = cerca_file_esistenti()
    file_ripreso = None
    nome_ripreso = None

    if file_esistenti:
        file_ripreso = chiedi_ripresa_o_nuovo(file_esistenti)

    if file_ripreso:
        wb = openpyxl.load_workbook(file_ripreso)
        nome_ripreso = file_ripreso.name
        print(f"\n  Ripreso: {file_ripreso}")
        prima_data = file_ripreso.stem
    else:
        wb = openpyxl.load_workbook(TEMPLATE_FILE)

    ws_riepilogo = wb["riepilogo_settimana"]
    ws_log = wb["dati_grezzi"]

    # Retrocompatibilita': rimuovi il vecchio foglio "bollettino" separato
    if "bollettino" in wb.sheetnames:
        del wb["bollettino"]

    display_menu()

    # Chiedi la settimana: se il file ripreso contiene gia' una settimana, proponila
    settimana_file = None
    if file_ripreso:
        settimana_file = leggi_settimana_da_file(ws_riepilogo)

    lunedi = chiedi_settimana(settimana_file)
    lunedi_str = lunedi.strftime("%d-%m-%Y")

    # Salva l'autosave se il processo riceve SIGTERM (es. chiusura GUI)
    if hasattr(signal, "SIGTERM"):
        def _sigterm_handler(signum, frame):
            try:
                _attendi_autosave()
                autosave(wb, lunedi_str, sincrono=True)
            except Exception:
                pass
            sys.exit(0)
        signal.signal(signal.SIGTERM, _sigterm_handler)

    if not file_ripreso:
        prima_data = lunedi_str

    # Aggiorna sempre l'intestazione (utile se l'utente ha cambiato settimana)
    compila_intestazione(ws_riepilogo, lunedi)

    log_row = find_next_log_row(ws_log)

    # Stato persistente tra sessioni giornaliere
    stato = {"ultimo_codice": None, "storico": [], "beep": False, "file_salvato": False}

    # Contesto invariante della sessione (passato a sessione_giorno)
    ctx = {
        "ws_riepilogo": ws_riepilogo, "ws_log": ws_log, "wb": wb,
        "lunedi": lunedi, "lunedi_str": lunedi_str,
        "prima_data": prima_data, "nome_ripreso": nome_ripreso,
        "file_ripreso": file_ripreso,
    }

    risultato = "quit"
    while True:
        giorno_num, data_str = chiedi_giorno(lunedi)

        # Protezione doppio giorno
        totale_esistente = controlla_giorno_esistente(ws_riepilogo, giorno_num)
        if totale_esistente > 0:
            nome_g = GIORNI_NOMI[giorno_num].upper()
            print(f"\n  ATTENZIONE: {nome_g} contiene gia' {totale_esistente} osservazioni.")
            risposta = input("  Continuare aggiungendo dati? (s/n): ").strip().lower()
            if risposta != "s":
                continue

        risultato, log_row = sessione_giorno(
            ctx, giorno_num, data_str, log_row, stato,
        )

        if risultato == "quit":
            break

        risposta = input("\nVuoi continuare con un altro giorno? (s/n): ").strip().lower()
        if risposta != "s":
            salvataggio_eseguito, cartella_salvata = menu_uscita_salvataggio(wb, prima_data, nome_ripreso, file_ripreso, stato)
            stato["file_salvato"] = salvataggio_eseguito
            stato["cartella_salvata"] = cartella_salvata
            break

    # Se file_salvato == True, il file è stato salvato dal menu di uscita
    # Se file_salvato == False, l'utente ha scelto di uscire senza salvare
    salvataggio_ok = stato.get("file_salvato", False)
    if salvataggio_ok and not file_ripreso:
        print("\n  Il template originale non e' stato modificato.")

    # Menu operazioni aggiuntive (nello stesso punto del codice originale)
    cartella_salvata = stato.get("cartella_salvata")
    if salvataggio_ok and cartella_salvata:
        print("\n  Operazioni aggiuntive:")
        print("    1. Aggiorna riepilogo annuale")
        print("    2. Genera bollettini Word (ITA/ENG)")
        print("    3. Entrambi (annuale + bollettini)")
        print("    Invio = nessuna operazione aggiuntiva")
        risp_extra = input("\n  Scelta: ").strip()
        if risp_extra in ("1", "3"):
            esporta_riepilogo_annuale(ws_riepilogo, lunedi, cartella_salvata)
        if risp_extra in ("2", "3"):
            genera_bollettini_word(ws_riepilogo, lunedi, lunedi_str, cartella_salvata)

    wb.close()
    pulisci_file_temporanei(lunedi_str, file_ripreso, salvataggio_ok)

    print("\n  Sessione terminata.\n")

    if getattr(sys, "frozen", False):
        input("Premi INVIO per chiudere...")


if __name__ == "__main__":
    main()
